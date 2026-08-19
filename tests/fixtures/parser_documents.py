"""Deterministic byte fixtures for document parser adapter tests."""

from datetime import date, datetime, time
from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from openpyxl import Workbook  # type: ignore[import-untyped]
from pypdf import PdfWriter


def synthetic_txt_bytes() -> bytes:
    """Return UTF-8 BOM-prefixed synthetic plain text."""
    return "\ufeff合成船厂规范\r\n\r\n泵组检查。\r\n第二行。".encode("utf-8")


def synthetic_markdown_bytes() -> bytes:
    """Return synthetic Markdown with headings, prose, and a table."""
    return (
        "# 合成规范\n\n"
        "## 泵组\n\n"
        "检查轴封。\n\n"
        "| 项目 | 数量 |\n| --- | ---: |\n| 泵 | 2 |\n"
    ).encode()


def synthetic_docx_bytes() -> bytes:
    """Return synthetic DOCX content with ordered semantic body blocks."""
    document = Document()
    document.add_heading("\u5408\u6210\u8239\u7ea7\u89c4\u5219", level=0)
    document.add_heading("\u673a\u68b0\u7cfb\u7edf", level=1)
    document.add_paragraph("\u68c0\u67e5\u6cf5\u7ec4\u8f74\u5c01\u3002")
    document.add_heading("\u6cf5\u7ec4", level=2)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "\u68c0\u67e5\u9879"
    table.cell(0, 1).text = "\u7ed3\u679c"
    table.cell(1, 0).text = "\u8f74\u5c01"
    table.cell(1, 1).text = "\u5408\u683c"

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def blank_docx_bytes() -> bytes:
    """Return a structurally valid DOCX without parseable body content."""
    document = Document()
    document.add_paragraph(" \t ")
    document.add_table(rows=1, cols=2)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def merged_table_docx_bytes() -> bytes:
    """Return a DOCX whose first table row contains a horizontal merge."""
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "\u68c0\u67e5\u9879"
    table.cell(1, 0).text = "\u8f74\u5c01"
    table.cell(1, 1).text = "\u5408\u683c"

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def heading_hierarchy_docx_bytes() -> bytes:
    """Return a DOCX covering title roots and exact heading style matching."""
    document = Document()
    document.add_heading("根标题", level=0)
    document.add_heading("初始系统", level=1)
    document.add_heading("初始设备", level=2)
    document.add_heading("替换系统", level=1)
    document.add_heading("九级主题", level=9)
    document.add_heading("八级替换", level=8)
    document.styles.add_style("Heading 10", WD_STYLE_TYPE.PARAGRAPH)
    document.add_paragraph("不是标题", style="Heading 10")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def synthetic_xlsx_bytes() -> bytes:
    """Return a synthetic workbook with ordered typed worksheet tables."""
    workbook = Workbook()
    workbook.iso_dates = True
    try:
        pumps = workbook.active
        pumps.title = "泵组"
        pumps.append(("项目", "计算", "日期", "通过", "备注", "状态"))
        pumps.append(("轴封", "=1+1", date(2026, 8, 19), True, None, "待复核"))
        pumps.append(
            (
                "泵轴",
                1.25,
                datetime(2026, 8, 19, 14, 30, 45),
                False,
                time(6, 15, 30),
                "完成",
            )
        )

        workbook.create_sheet("空白")
        materials = workbook.create_sheet("材料")
        materials.append(("材料", "数量"))
        materials.append(("钢板", 12))

        output = BytesIO()
        workbook.save(output)
        return output.getvalue()
    finally:
        workbook.close()


def blank_xlsx_bytes() -> bytes:
    """Return a workbook containing only blank worksheets."""
    workbook = Workbook()
    try:
        workbook.active.title = "空白一"
        workbook.active.append((None, None))
        workbook.create_sheet("空白二")
        output = BytesIO()
        workbook.save(output)
        return output.getvalue()
    finally:
        workbook.close()


def trailing_blank_xlsx_bytes() -> bytes:
    """Return a workbook with interior and trailing blank table cells."""
    workbook = Workbook()
    try:
        worksheet = workbook.active
        worksheet.title = "裁尾"
        worksheet.append(("项目", "数量", "备注", ""))
        worksheet.append(("泵", 2, None, ""))
        worksheet.append(("", "", "", ""))
        output = BytesIO()
        workbook.save(output)
        return output.getvalue()
    finally:
        workbook.close()


def synthetic_pdf_bytes() -> bytes:
    """Return a deterministic two-page PDF with extractable text streams."""
    return _minimal_pdf_bytes(("Synthetic page one", "Synthetic page two"))


def blank_pdf_bytes() -> bytes:
    """Return a valid one-page PDF without a text layer."""
    writer = PdfWriter()
    try:
        writer.add_blank_page(width=612, height=792)
        output = BytesIO()
        writer.write(output)
        return output.getvalue()
    finally:
        writer.close()


def encrypted_pdf_bytes() -> bytes:
    """Return a valid synthetic password-protected PDF."""
    writer = PdfWriter()
    try:
        writer.add_blank_page(width=612, height=792)
        writer.encrypt("synthetic-test-password")
        output = BytesIO()
        writer.write(output)
        return output.getvalue()
    finally:
        writer.close()


def pdf_with_blank_middle_page_bytes() -> bytes:
    """Return text pages separated by one valid blank page."""
    return _minimal_pdf_bytes(("First text page", None, "Third text page"))


def _minimal_pdf_bytes(page_texts: tuple[str | None, ...]) -> bytes:
    object_bodies: list[bytes] = []
    page_references: list[bytes] = []
    font_object_number = 3 + 2 * len(page_texts)

    for page_index, page_text in enumerate(page_texts):
        page_object_number = 3 + 2 * page_index
        content_object_number = page_object_number + 1
        page_references.append(f"{page_object_number} 0 R".encode())
        resources = (
            f"/Resources << /Font << /F1 {font_object_number} 0 R >> >> "
            if page_text is not None
            else ""
        )
        object_bodies.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"{resources}/Contents {content_object_number} 0 R >>"
            ).encode()
        )
        stream = (
            f"BT /F1 12 Tf 72 720 Td ({page_text}) Tj ET".encode()
            if page_text is not None
            else b""
        )
        object_bodies.append(
            b"<< /Length "
            + str(len(stream)).encode()
            + b">>\nstream\n"
            + stream
            + b"\nendstream"
        )

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids ["
        + b" ".join(page_references)
        + b"] /Count "
        + str(len(page_texts)).encode()
        + b" >>",
        *object_bodies,
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    output = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for object_number, body in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{object_number} 0 obj\n".encode())
        output.extend(body)
        output.extend(b"\nendobj\n")

    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode())
    output.extend(
        b"trailer\n<< /Size "
        + str(len(objects) + 1).encode()
        + b" /Root 1 0 R >>\nstartxref\n"
        + str(xref_offset).encode()
        + b"\n%%EOF\n"
    )
    return bytes(output)
